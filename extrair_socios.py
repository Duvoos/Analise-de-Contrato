
                    ####### VERSÃO OTIMIZADA #########


import docx  # Biblioteca usada para manipular arquivos do Word (.docx)
import re  # Biblioteca de expressões regulares para buscar padrões de texto

# Função para ler o documento .docx e obter o texto
def read_docx(doc_path):
    """
    Lê o conteúdo de um arquivo .docx e retorna o texto completo.
    
    :param doc_path: Caminho para o arquivo .docx
    :return: Texto do documento como uma única string
    """
    doc = docx.Document(doc_path)  # Abre o documento .docx
    full_text = []
    
    # Itera por cada parágrafo do documento e adiciona o texto à lista full_text
    for para in doc.paragraphs:
        full_text.append(para.text.strip())  # Remove espaços desnecessários
    return "\n".join(full_text)  # Retorna o texto completo como uma string única

# Função para extrair sócios e suas respectivas cotas do texto
def extract_socios_cotas(text):
    """
    Extrai os nomes dos sócios e suas cotas a partir do texto.
    
    :param text: Texto contendo as informações do contrato
    :return: Lista de dicionários com os nomes dos sócios e a quantidade de cotas
    """
    socios = []  # Lista para armazenar os dados dos sócios

    # Expressão regular para identificar números seguidos da palavra "cotas"
    cotas_pattern = re.compile(r'(\d+)\s+cotas')

    # Divide o texto em linhas e itera sobre cada linha
    for line in text.split('\n'):
        # Verifica se a linha contém a palavra "cotas"
        if "cotas" in line.lower():
            # Extrai o nome do sócio (texto antes da palavra "portador")
            name = line.split('portador')[0].strip()
            
            # Busca a quantidade de cotas usando a expressão regular
            cotas_match = cotas_pattern.search(line)
            if cotas_match:
                cotas = int(cotas_match.group(1))  # Extrai o número de cotas
                # Adiciona o nome do sócio e a quantidade de cotas à lista 'socios'
                socios.append({"Sócio": name, "Cotas": cotas})

    return socios  # Retorna a lista de sócios e suas cotas

# Função para gerar um arquivo .docx com os dados extraídos
def gerar_docx(socios, output_path):
    """
    Gera um arquivo .docx com o quadro societário extraído.

    :param socios: Lista de sócios e suas cotas
    :param output_path: Caminho para salvar o arquivo gerado
    """
    doc = docx.Document()  # Cria um novo documento Word
    doc.add_heading('Quadro Societário', level=1)  # Adiciona o título "Quadro Societário"

    # Cria uma tabela com duas colunas: "Sócio" e "Cotas"
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sócio'
    hdr_cells[1].text = 'Cotas'

    # Preenche a tabela com os dados dos sócios
    for socio in socios:
        row_cells = table.add_row().cells
        row_cells[0].text = socio["Sócio"]  # Nome do sócio
        row_cells[1].text = str(socio["Cotas"])  # Quantidade de cotas

    # Salva o documento gerado no caminho especificado
    doc.save(output_path)

# Função principal que executa o processo completo de leitura, extração e geração do arquivo
if __name__ == "__main__":
    # Define o caminho do arquivo de entrada (o contrato de partnership)
    doc_path = r"C:\Users\Eduardo\Desktop\Etapa02\Partnership.docx"

    # Define o caminho do arquivo de saída (o documento com o quadro societário gerado)
    output_path = r"C:\Users\Eduardo\Desktop\Etapa02\quadro_societario.docx"

    # Lê o arquivo .docx e obtém o texto completo do contrato
    text = read_docx(doc_path)

    # Extrai os sócios e suas cotas a partir do texto
    socios = extract_socios_cotas(text)

    # Gera o arquivo .docx com os dados extraídos (nomes dos sócios e suas cotas)
    gerar_docx(socios, output_path)

    # Exibe uma mensagem de sucesso ao final do processo
    print(f"Arquivo gerado com sucesso: {output_path}")


                      ###### VERSÃO SIMPLIFICADA ######

# import docx  # Biblioteca usada para manipular arquivos do Word (.docx)

# # Função para ler o documento .docx e obter o texto
# def read_docx(doc_path):
#     """
#     Lê o conteúdo de um arquivo .docx e retorna o texto completo.
#     """
#     doc = docx.Document(doc_path)  # Abre o documento .docx
#     full_text = []
    
#     # Itera por cada parágrafo e armazena o texto na lista full_text
#     for para in doc.paragraphs:
#         full_text.append(para.text.strip())  # Remove espaços extras e armazena o texto
#     return "\n".join(full_text)  # Junta o texto de todos os parágrafos em uma string única

# # Função para extrair os nomes dos sócios e suas cotas
# def extract_socios_cotas(text):
#     """
#     Extrai os nomes dos sócios e suas cotas a partir do texto.
#     """
#     socios = []  # Lista para armazenar os sócios e cotas
    
#     # Divide o texto em linhas e processa cada uma
#     lines = text.split("\n")
#     for line in lines:
#         if "cotas" in line:  # Verifica se a linha contém a palavra "cotas"
#             try:
#                 # Extrai o nome e as cotas de forma mais manual
#                 nome = line.split(",")[0]
#                 cotas = int(line.split("cotas")[0].split()[-1])
#                 socios.append((nome, cotas))
#             except:
#                 pass  # Simplesmente ignora linhas que não estão no formato esperado
    
#     return socios  # Retorna uma lista de tuplas (nome, cotas)

# # Função para gerar um arquivo de saída com os dados extraídos
# def gerar_docx(socios, output_path):
#     """
#     Gera um arquivo .docx com o quadro societário extraído.
#     """
#     doc = docx.Document()  # Cria um novo documento Word
#     doc.add_heading('Quadro Societário', level=1)  # Adiciona um título ao documento

#     # Para cada sócio, cria uma linha no documento com o nome e as cotas
#     for socio in socios:
#         nome, cotas = socio
#         doc.add_paragraph(f"{nome}: {cotas} cotas")

#     # Salva o arquivo gerado
#     doc.save(output_path)

# # Parte principal do código
# if __name__ == "__main__":
#     # Definir o caminho do arquivo de entrada e de saída
#     doc_path = "Partnership.docx"  # Arquivo de entrada
#     output_path = "quadro_societario.docx"  # Arquivo de saída

#     # Ler o conteúdo do arquivo .docx
#     texto = read_docx(doc_path)

#     # Extrair sócios e cotas do texto
#     socios = extract_socios_cotas(texto)

#     # Gerar o arquivo .docx com o resultado
#     gerar_docx(socios, output_path)

#     # Mensagem de sucesso
#     print(f"Arquivo gerado com sucesso: {output_path}")
